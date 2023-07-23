""" 
    Take an unformatted quiz document and properly format it with spacing, indentation, and numbering.
    Use cases: Generating a quiz using an AI chatbot for studying purposes, and running this script to automatically format it

    Dev: Emily Rose
    Last Updated: July 22, 2023
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK

letterOptions = ['a', 'b', 'c', 'd', 'e', 'f', 'g']
fileName = input("What is the name of the document you wish to format (case sensitive)? Please exclude the '.docx' in your input.")
directory = input("Please input the absolute file path of your document.")

# Read and open the requested file at the specified location
os.chdir(f"{directory}")
wb = Document(rf"{fileName}"+".docx")

# Normalise all styling in the document to unstylised Arial 12 pt font
def styleNormalise(paragraph):
    for run in paragraph.runs:
        run.font.size = Pt(12) 
        run.font.name = "Arial"
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False

# Add a new paragraph above the question. If there is more text following the question, split
# it from the question itself. Bold the question.
def setQuestion(run, start_index):
    paragraph.insert_paragraph_before().runs.append(paragraph.add_run(run.text[start_index:]))
    run.text = run.text[:start_index] + "\n"
    run.font.bold = True

# Split the question apart from the answers:
# Find the end of the question, then add a new paragraph before the question that only states the
# question. Then, edit the current run to contain only the answer options
# Bold all questions and add space beneath them
def questionNormalise(paragraph):
    for run in paragraph.runs:

        # Determine what is a question and what is not
        if paragraph.text.endswith(":"):                # Question
            start_index = run.text.find(":") + 1
            setQuestion(run, start_index)
        elif "?" in run.text:                           # Question
            start_index = run.text.find("?") + 1
            setQuestion(run, start_index)
        elif not run.text.isspace():                    # Other
            start_index = 0
            paragraph.add_run(run.text)
            run.text = ""
            break
        else: 
            break

        # Iterate over the newly edited paragraph runs
        i = 0
        for run in paragraph.runs:
            if re.search("[a-zA-Z]\.", run.text):               # Look for a question answer in the format "a."
                questions = re.split("[a-zA-Z]\.", run.text)    # Create a list separated by answers
                run.text = ""

                # Empty list of valid questions (ie. no questions that are just whitespace or commas)
                fixedQuestions = []
                for question in questions:                      # Search for valid questions and append them to the new list
                    if question.strip():
                        fixedQuestions.append(question)
                
                for question in fixedQuestions:                  # Write the valid questions onto the doc
                    paragraph.add_run(letterOptions[i] + ". " + question + "\n")
                    if i < len(questions): i = i + 1            # Prefix the question with a letter in letterOptions

# Indent multiple choice questions
def optionsNormalise(paragraph):
    if re.search("[a-zA-Z]\)|[a-zA-Z]\.", paragraph.text) != None:  # Search for all multiple choice options
        paragraph.paragraph_format.left_indent = Inches(0.5)        # Indent all multiple choice questions

# Break the answers page onto a new page
def answerNormalise(paragraph):
    if "answer" in paragraph.text.lower():                      # Look for the answers page
        paragraph.add_run(paragraph.runs[0].text)               # Add a new run to the end of the paragraph with the same text as the first
        paragraph.runs[0].text = ""                             # Clear the text at the start of the run
        paragraph.runs[0].add_break(WD_BREAK.PAGE)              # Add a page break where the start of the answers page used to be

# Format the entire document, paragraph by paragraph
for paragraph in wb.paragraphs:
    styleNormalise(paragraph)
    if paragraph.text.endswith(":") or "?" in paragraph.text:
        questionNormalise(paragraph)
    optionsNormalise(paragraph)
    answerNormalise(paragraph)

wb.save(f"{fileName}_formatted.docx")                           # Save a new version of the document